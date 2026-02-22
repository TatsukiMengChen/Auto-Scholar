import io
from enum import StrEnum

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from backend.schemas import CitationStyle, DraftOutput, PaperMetadata


class ExportFormat(StrEnum):
    MARKDOWN = "markdown"
    DOCX = "docx"


def _format_authors_apa(authors: list[str]) -> str:
    if not authors:
        return ""
    if len(authors) == 1:
        return authors[0]
    if len(authors) == 2:
        return f"{authors[0]} & {authors[1]}"
    if len(authors) <= 20:
        return ", ".join(authors[:-1]) + ", & " + authors[-1]
    return ", ".join(authors[:19]) + ", ... " + authors[-1]


def _format_authors_mla(authors: list[str]) -> str:
    if not authors:
        return ""
    if len(authors) == 1:
        return authors[0]
    if len(authors) == 2:
        return f"{authors[0]}, and {authors[1]}"
    return f"{authors[0]}, et al."


def _format_authors_ieee(authors: list[str]) -> str:
    if not authors:
        return ""
    if len(authors) <= 3:
        return ", ".join(authors)
    return f"{authors[0]} et al."


def _format_authors_gbt7714(authors: list[str]) -> str:
    if not authors:
        return ""
    if len(authors) <= 3:
        return ", ".join(authors)
    return ", ".join(authors[:3]) + ", 等"


def format_citation(paper: PaperMetadata, index: int, style: CitationStyle) -> str:
    authors = paper.authors
    title = paper.title
    year = paper.year
    url = paper.url

    if style == CitationStyle.APA:
        author_str = _format_authors_apa(authors)
        year_str = f"({year})" if year else "(n.d.)"
        ref = f"{author_str} {year_str}. {title}."
        if url:
            ref += f" {url}"
        return ref

    elif style == CitationStyle.MLA:
        author_str = _format_authors_mla(authors)
        year_str = str(year) if year else "n.d."
        ref = f'{author_str}. "{title}." {year_str}.'
        if url:
            ref += f" {url}"
        return ref

    elif style == CitationStyle.IEEE:
        author_str = _format_authors_ieee(authors)
        year_str = str(year) if year else "n.d."
        ref = f'[{index}] {author_str}, "{title}," {year_str}.'
        if url:
            ref += f" [Online]. Available: {url}"
        return ref

    elif style == CitationStyle.GB_T7714:
        author_str = _format_authors_gbt7714(authors)
        year_str = str(year) if year else ""
        ref = f"[{index}] {author_str}. {title}[J]. {year_str}."
        if url:
            ref += f" {url}"
        return ref

    return f"{index}. {title}"


def format_references(papers: list[PaperMetadata], style: CitationStyle) -> list[str]:
    return [format_citation(p, i, style) for i, p in enumerate(papers, 1)]


def export_to_markdown(
    draft: DraftOutput,
    papers: list[PaperMetadata],
    citation_style: CitationStyle = CitationStyle.APA,
) -> str:
    lines: list[str] = []
    lines.append(f"# {draft.title}\n")

    for section in draft.sections:
        lines.append(f"## {section.heading}\n")
        lines.append(f"{section.content}\n")

    if papers:
        lines.append("## References\n")
        refs = format_references(papers, citation_style)
        for ref in refs:
            lines.append(f"{ref}\n")

    return "\n".join(lines)


def export_to_docx(
    draft: DraftOutput,
    papers: list[PaperMetadata],
    citation_style: CitationStyle = CitationStyle.APA,
) -> bytes:
    doc = Document()

    title = doc.add_heading(draft.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in draft.sections:
        doc.add_heading(section.heading, level=1)
        para = doc.add_paragraph(section.content)
        para.paragraph_format.first_line_indent = Inches(0.5)

    if papers:
        doc.add_heading("References", level=1)
        refs = format_references(papers, citation_style)
        for ref in refs:
            ref_para = doc.add_paragraph(ref)
            ref_para.paragraph_format.space_after = Pt(6)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
